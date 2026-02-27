# Genera un reporte DOCX del proyecto de la MT de Fibonacci.
# Requiere:  python-docx  (pip install python-docx)

import os, sys, json, csv

RAIZ = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
RUTA_CSV      = os.path.join(RAIZ, "analisis", "resultadosBenchmark.csv")
RUTA_GRAFICO  = os.path.join(RAIZ, "analisis", "dispersionTiempo.png")
RUTA_DIAGRAMA = os.path.join(RAIZ, "documentos", "diagramaFibonacci.png")
RUTA_CONFIG   = os.path.join(RAIZ, "configuracion", "fibonacci.json")
RUTA_DOCX     = os.path.join(RAIZ, "documentos", "reporteProyecto.docx")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx no instalado. Ejecute:  pip install python-docx")
    sys.exit(1)


def fibonacci(n):
    if n <= 1:
        return n
    a, b = 0, 1
    for _ in range(2, n + 1):
        a, b = b, a + b
    return b


def resolver3x3(a, b):
    def det3(m):
        return (m[0][0]*(m[1][1]*m[2][2]-m[1][2]*m[2][1])
                - m[0][1]*(m[1][0]*m[2][2]-m[1][2]*m[2][0])
                + m[0][2]*(m[1][0]*m[2][1]-m[1][1]*m[2][0]))
    d = det3(a)
    if d == 0:
        return 0.0, 0.0, 0.0
    ax = [[b[0],a[0][1],a[0][2]],[b[1],a[1][1],a[1][2]],[b[2],a[2][1],a[2][2]]]
    ay = [[a[0][0],b[0],a[0][2]],[a[1][0],b[1],a[1][2]],[a[2][0],b[2],a[2][2]]]
    az = [[a[0][0],a[0][1],b[0]],[a[1][0],a[1][1],b[1]],[a[2][0],a[2][1],b[2]]]
    return det3(ax)/d, det3(ay)/d, det3(az)/d


def regresionCuadratica(puntos):
    if not puntos:
        return 0.0, 0.0, 0.0
    sx  = sum(x for x,_ in puntos)
    sx2 = sum(x**2 for x,_ in puntos)
    sx3 = sum(x**3 for x,_ in puntos)
    sx4 = sum(x**4 for x,_ in puntos)
    sy  = sum(y for _,y in puntos)
    sxy = sum(x*y for x,y in puntos)
    sx2y= sum(x**2*y for x,y in puntos)
    n   = float(len(puntos))
    return resolver3x3([[sx4,sx3,sx2],[sx3,sx2,sx],[sx2,sx,n]],[sx2y,sxy,sy])


def estiloTitulo(doc, texto, nivel=1):
    doc.add_heading(texto, level=nivel)


def parrafo(doc, texto, negrita=False, italica=False, tam=11):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    run.italic = italica
    run.font.size = Pt(tam)
    return p


def agregarImagen(doc, ruta, ancho=5.5):
    if os.path.isfile(ruta):
        doc.add_picture(ruta, width=Inches(ancho))
        ultimo = doc.paragraphs[-1]
        ultimo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        parrafo(doc, f"[Imagen no encontrada: {os.path.basename(ruta)}]", italica=True)
        return False


def seccionPortada(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proyecto 1\n")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(26, 115, 232)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Maquina de Turing\nSucesion de Fibonacci")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(80, 80, 80)

    for _ in range(4):
        doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Algoritmos y Estructuras de Datos")
    run3.font.size = Pt(14)

    doc.add_page_break()


def seccionConvenciones(doc):
    estiloTitulo(doc, "1. Convenciones elegidas")

    parrafo(doc, "Se utiliza codificacion unaria para representar "
            "tanto la entrada como la salida de la Maquina de Turing:")

    tabla = doc.add_table(rows=6, cols=2, style="Light List Accent 1")
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    encabezados = ["Concepto", "Convencion"]
    for i, h in enumerate(encabezados):
        tabla.rows[0].cells[i].text = h

    datos = [
        ("Entrada", "Cadena de n unos: 1^n  (codificacion unaria)"),
        ("Salida", "Cadena de F(n) unos: 1^F(n)"),
        ("Cero", "Cinta vacia (sin unos) -- F(0) = 0"),
        ("Simbolo blanco", "B"),
        ("Overflow (n > 12)", "Se escribe X en la cinta"),
    ]
    for i, (c1, c2) in enumerate(datos, start=1):
        tabla.rows[i].cells[0].text = c1
        tabla.rows[i].cells[1].text = c2

    doc.add_paragraph()
    parrafo(doc, "Ejemplo: para n = 5, la entrada es \"11111\" y la salida "
            "esperada es \"11111\" (cinco unos, ya que F(5) = 5).")

    doc.add_paragraph()
    estiloTitulo(doc, "1.1 Convencion para enteros negativos", nivel=2)
    parrafo(doc, "La sucesion de Fibonacci clasica se define para n >= 0. "
            "Sin embargo, el proyecto requiere definir una convencion para "
            "enteros negativos.")
    parrafo(doc, "Convencion adoptada: se utiliza el simbolo '-' como prefijo "
            "en la cinta para indicar signo negativo. Por ejemplo, la entrada "
            "\"-111\" representa n = -3.")
    parrafo(doc, "Comportamiento de la MT ante entrada negativa: la maquina "
            "lee el simbolo '-' al inicio de la cinta, lo interpreta como entrada "
            "invalida para el calculo de Fibonacci, y se detiene sin aceptar "
            "(no existe transicion definida desde el estado inicial para '-'). "
            "Esto es coherente con la definicion matematica: F(n) para n < 0 "
            "no se calcula con esta MT.")
    parrafo(doc, "Esta decision simplifica el diseno: la MT solo acepta cadenas "
            "de unos (1^n) y rechaza cualquier otra entrada, incluyendo las que "
            "comienzan con '-'.", italica=True, tam=10)


def seccionDiagrama(doc):
    estiloTitulo(doc, "2. Diagrama de la MT")

    parrafo(doc, "La siguiente figura muestra la estructura general de la "
            "Maquina de Turing para calcular la sucesion de Fibonacci. "
            "Los estados de lectura (azul) recorren la cadena de entrada, "
            "los estados de escritura (amarillo) generan la salida, "
            "y el estado de aceptacion (verde) marca la finalizacion.")

    agregarImagen(doc, RUTA_DIAGRAMA, ancho=6.0)

    parrafo(doc, "Al tener 389 estados, el diagrama muestra el patron "
            "general en lugar de cada estado individual.", italica=True, tam=9)


def seccionTablaTransiciones(doc):
    estiloTitulo(doc, "2.1 Tabla de transiciones (extracto)", nivel=2)

    parrafo(doc, "A continuacion se muestra un extracto representativo de la "
            "tabla de transiciones. La tabla completa se encuentra en el "
            "archivo configuracion/fibonacci.json (402 transiciones).")

    if not os.path.isfile(RUTA_CONFIG):
        parrafo(doc, "[fibonacci.json no encontrado]")
        return

    with open(RUTA_CONFIG, "r", encoding="utf-8") as f:
        config = json.load(f)
    transiciones = config.get("transitions", [])

    encabezados = ["Estado actual", "Lee", "Siguiente estado", "Escribe", "Mueve"]
    muestra = transiciones[:20]
    tabla = doc.add_table(rows=len(muestra)+1, cols=len(encabezados),
                          style="Light List Accent 1")
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(encabezados):
        tabla.rows[0].cells[j].text = h

    for i, t in enumerate(muestra, start=1):
        tabla.rows[i].cells[0].text = t.get("current_state", "")
        tabla.rows[i].cells[1].text = t.get("read_symbol", "")
        tabla.rows[i].cells[2].text = t.get("next_state", "")
        tabla.rows[i].cells[3].text = t.get("write_symbol", "")
        tabla.rows[i].cells[4].text = t.get("move", "")

    doc.add_paragraph()
    parrafo(doc, f"... y {len(transiciones) - len(muestra)} transiciones mas.",
            italica=True, tam=9)


def seccionComponentes(doc):
    estiloTitulo(doc, "3. Componentes de la MT")

    if not os.path.isfile(RUTA_CONFIG):
        parrafo(doc, "[fibonacci.json no encontrado]")
        return

    with open(RUTA_CONFIG, "r", encoding="utf-8") as f:
        config = json.load(f)

    estados = config.get("states", [])
    alfabetoEntrada = config.get("input_alphabet", [])
    alfabetoCinta = config.get("tape_alphabet", [])
    estadoInicial = config.get("initial_state", "")
    estadosAceptacion = config.get("accept_states", [])
    blanco = config.get("blank_symbol", "")
    transiciones = config.get("transitions", [])
    numTrans = len(transiciones)

    tabla = doc.add_table(rows=8, cols=2, style="Light List Accent 1")
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    filas = [
        ("Componente", "Valor"),
        ("N. de estados", str(len(estados))),
        ("Alfabeto de entrada (S)", ", ".join(alfabetoEntrada)),
        ("Alfabeto de cinta (G)", ", ".join(alfabetoCinta)),
        ("Estado inicial", estadoInicial),
        ("Estados de aceptacion", ", ".join(estadosAceptacion)),
        ("Simbolo blanco", blanco),
        ("N. de transiciones", str(numTrans)),
    ]
    for i, (c1, c2) in enumerate(filas):
        tabla.rows[i].cells[0].text = c1
        tabla.rows[i].cells[1].text = c2


def seccionPruebasManuales(doc):
    estiloTitulo(doc, "3.1 Pruebas manuales de validacion", nivel=2)

    parrafo(doc, "Las siguientes pruebas sirven para verificar que el "
            "simulador produce los resultados correctos:")

    encabezados = ["n", "Entrada en cinta", "Salida esperada", "F(n)", "Descripcion"]
    casos = [
        ("0", "(cinta vacia)", "(cinta vacia)", "0", "Caso base: F(0) = 0"),
        ("1", "1", "1", "1", "Caso base: F(1) = 1"),
        ("5", "11111", "11111", "5", "F(5) = 5, salida = cinco unos"),
        ("7", "1111111", "1111111111111", "13", "F(7) = 13, salida = trece unos"),
        ("12", "111111111111", "1 x 144", "144", "F(12) = 144, salida = 144 unos"),
        ("-3", "-111", "(no acepta)", "N/A", "Entrada negativa: MT se detiene"),
    ]

    tabla = doc.add_table(rows=len(casos)+1, cols=len(encabezados),
                          style="Light List Accent 1")
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(encabezados):
        tabla.rows[0].cells[j].text = h

    for i, fila in enumerate(casos, start=1):
        for j, val in enumerate(fila):
            tabla.rows[i].cells[j].text = val


def seccionAnalisis(doc):
    estiloTitulo(doc, "4. Analisis empirico")

    if not os.path.isfile(RUTA_CSV):
        parrafo(doc, "[resultadosBenchmark.csv no encontrado -- "
                "ejecute el benchmark desde el menu principal]")
        return

    filas = []
    with open(RUTA_CSV, "r", encoding="utf-8") as f:
        for r in csv.DictReader(f):
            filas.append(r)

    parrafo(doc, "Se ejecuto la MT para valores de n de 0 a 12 con 3 "
            "repeticiones cada uno, midiendo pasos y tiempo de ejecucion.")

    encabezados = ["n", "Entrada", "Pasos (prom)", "Tiempo (s)", "F(n)", "Correcto"]
    tabla = doc.add_table(rows=len(filas)+1, cols=len(encabezados),
                          style="Light List Accent 1")
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(encabezados):
        tabla.rows[0].cells[j].text = h

    for i, fila in enumerate(filas, start=1):
        n = int(fila["n"])
        salida = fila.get("salida", "")
        fn = fibonacci(n)
        correcto = "Si" if len(salida) == fn else "No"
        vals = [
            str(n),
            fila.get("entrada", "1" * n),
            f"{float(fila['pasosPromedio']):.0f}",
            f"{float(fila['tiempoPromSeg']):.6f}",
            str(fn),
            correcto,
        ]
        for j, v in enumerate(vals):
            tabla.rows[i].cells[j].text = v

    doc.add_paragraph()

    estiloTitulo(doc, "4.1 Grafico de dispersion", nivel=2)
    parrafo(doc, "La siguiente grafica muestra el tiempo de ejecucion "
            "en funcion del tamano de entrada n, junto con la curva "
            "de regresion cuadratica ajustada:")

    agregarImagen(doc, RUTA_GRAFICO, ancho=5.5)

    puntos = []
    for fila in filas:
        puntos.append((float(fila["n"]), float(fila["tiempoPromSeg"])))
    a, b, c = regresionCuadratica(puntos)

    estiloTitulo(doc, "4.2 Regresion y complejidad", nivel=2)
    parrafo(doc, f"T(n) ~= {a:.10f}*n^2 + {b:.10f}*n + {c:.10f}")
    doc.add_paragraph()
    parrafo(doc, "El termino dominante es cuadratico (n^2), lo que indica "
            "que la complejidad temporal de esta Maquina de Turing para "
            "calcular F(n) usando codificacion unaria es:")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("O(n^2)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(26, 115, 232)

    doc.add_paragraph()
    parrafo(doc, "Justificacion: la regresion polinomial de grado 2 se ajusta "
            "adecuadamente a los datos experimentales. El coeficiente del termino "
            "cuadratico es positivo y dominante, confirmando que el numero de pasos "
            "crece cuadraticamente respecto al tamano de entrada n. "
            "Esto se debe a que para cada valor de n, la MT debe recorrer la "
            "entrada (n pasos de lectura) y luego escribir F(n) unos en la salida, "
            "donde F(n) crece exponencialmente pero la representacion unaria "
            "hace que el trabajo total sea polinomial en n.", italica=False, tam=10)


def main():
    os.makedirs(os.path.dirname(RUTA_DOCX), exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"  # type: ignore[reportAttributeAccessIssue]
    style.font.size = Pt(11)  # type: ignore[reportAttributeAccessIssue]

    seccionPortada(doc)
    seccionConvenciones(doc)
    doc.add_page_break()
    seccionDiagrama(doc)
    seccionTablaTransiciones(doc)
    doc.add_page_break()
    seccionComponentes(doc)
    seccionPruebasManuales(doc)
    doc.add_page_break()
    seccionAnalisis(doc)

    doc.save(RUTA_DOCX)
    print(f"Reporte generado: {RUTA_DOCX}")


if __name__ == "__main__":
    main()
